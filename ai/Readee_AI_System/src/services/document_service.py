"""
Document processing service cho Readee_AI_System.

Pipeline:
- PDF (Tối ưu tốc độ):
  + Thử extract text trực tiếp từ PDF trước (nhanh, ~1-2s cho file nhiều trang)
  + Nếu text đủ tốt (>=70% trang có text) → dùng text trực tiếp, không cần OCR
  + Nếu không đủ → OCR những trang thiếu text hoặc OCR toàn bộ
  + Trích ảnh từ PDF để moderation ảnh.
- DOCX:
  + Trích text + ảnh trực tiếp bằng python-docx.
  + Ghi text ra file tạm -> text_path.

Sau đó:
- Chia text thành nhiều đoạn nhỏ -> moderation text.
- Trả về (full_text, text_path, images) cho router dùng với moderation + summary.
"""

from __future__ import annotations

import io
import tempfile
import asyncio
from pathlib import Path
from typing import List, Tuple, Dict

import fitz  # PyMuPDF
from docx import Document
from PIL import Image

from src.services.ocr_client import run_ocr_on_file, run_ocr_on_file_parallel, run_ocr_on_pages


class DocumentService:
    def __init__(self) -> None:
        self.temp_dir = Path(tempfile.gettempdir())

    # --------------------- PDF helpers ---------------------

    def _extract_images_from_pdf(self, pdf_path: str) -> Tuple[List[Image.Image], List[int]]:
        """
        Trích ảnh từ PDF và trả về (images, page_numbers).
        page_numbers[i] = số trang (1-indexed) của images[i].
        Tối ưu: resize và convert ngay trong quá trình extract để tiết kiệm memory.
        """
        images: List[Image.Image] = []
        page_numbers: List[int] = []
        doc = fitz.open(pdf_path)
        max_size = (1280, 1280)  # Max 1280x1280 - đủ cho moderation, nhanh hơn
        
        try:
            for page_index in range(len(doc)):
                page = doc.load_page(page_index)
                page_num = page_index + 1  # 1-indexed
                
                for img in page.get_images():
                    xref = img[0]
                    base_image = doc.extract_image(xref)
                    image_bytes = base_image["image"]
                    
                    # Mở và xử lý ảnh - load toàn bộ vào memory trước khi đóng buffer
                    img_buffer = io.BytesIO(image_bytes)
                    img_obj = Image.open(img_buffer)
                    # Load toàn bộ image vào memory để tránh lỗi khi buffer đóng
                    img_obj.load()
                    # Đóng buffer sau khi đã load xong
                    img_buffer.close()
                    
                    # Convert sang RGB nếu cần (giảm memory cho RGBA)
                    if img_obj.mode != "RGB":
                        img_obj = img_obj.convert("RGB")
                    
                    # Resize nếu quá lớn (giảm RAM usage và tăng tốc độ inference)
                    if img_obj.size[0] > max_size[0] or img_obj.size[1] > max_size[1]:
                        img_obj.thumbnail(max_size, Image.Resampling.LANCZOS)
                    
                    images.append(img_obj)
                    page_numbers.append(page_num)
                    
                    # Clear image_bytes ngay sau khi tạo Image object
                    del image_bytes, base_image
                
                # Clear page object sau mỗi page để giải phóng memory
                del page
        finally:
            doc.close()
            del doc
        return images, page_numbers
    
    def _extract_text_by_pages_from_pdf(self, pdf_path: str) -> Tuple[str, Dict[int, Tuple[int, int]]]:
        """
        Trích text từ PDF theo từng trang và trả về:
        - full_text: toàn bộ text đã nối
        - page_ranges: dict mapping page_num -> (start_char, end_char) trong full_text
        """
        doc = fitz.open(pdf_path)
        page_texts: List[str] = []
        page_ranges: Dict[int, Tuple[int, int]] = {}
        
        try:
            current_pos = 0
            for page_index in range(len(doc)):
                page = doc.load_page(page_index)
                page_num = page_index + 1  # 1-indexed
                page_text = page.get_text()
                
                start_pos = current_pos
                end_pos = current_pos + len(page_text)
                
                # Nếu không phải trang cuối, thêm newline
                if page_index < len(doc) - 1:
                    page_text += "\n"
                    end_pos += 1
                
                page_texts.append(page_text)
                page_ranges[page_num] = (start_pos, end_pos)
                current_pos = end_pos
                
                del page
        finally:
            doc.close()
            del doc
        
        full_text = "".join(page_texts)
        del page_texts
        return full_text, page_ranges

    # --------------------- DOCX helpers ---------------------

    def _extract_text_and_images_from_docx(
        self, docx_path: str
    ) -> Tuple[str, List[Image.Image], List[int], Dict[int, Tuple[int, int]]]:
        """
        Trích text và ảnh từ DOCX.
        Trả về: (full_text, images, image_page_numbers, page_ranges)
        Note: DOCX không có khái niệm "trang" rõ ràng, nên ta ước tính dựa trên số đoạn.
        """
        doc = Document(docx_path)

        texts: List[str] = []
        for p in doc.paragraphs:
            if p.text.strip():
                texts.append(p.text.strip())

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        texts.append(cell.text.strip())

        images: List[Image.Image] = []
        image_page_numbers: List[int] = []
        
        # Ước tính page number cho ảnh dựa trên vị trí trong document
        # Giả sử mỗi ~50 dòng text = 1 trang (ước tính)
        paragraphs_per_page = 50
        current_paragraph_count = 0
        max_size = (1280, 1280)  # Max 1280x1280 - đủ cho moderation, nhanh hơn
        
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                img_data = rel.target_part.blob
                
                # Mở và xử lý ảnh - load toàn bộ vào memory trước khi đóng buffer
                img_buffer = io.BytesIO(img_data)
                img_obj = Image.open(img_buffer)
                # Load toàn bộ image vào memory để tránh lỗi khi buffer đóng
                img_obj.load()
                # Đóng buffer sau khi đã load xong
                img_buffer.close()
                
                # Convert sang RGB nếu cần (giảm memory cho RGBA)
                if img_obj.mode != "RGB":
                    img_obj = img_obj.convert("RGB")
                
                # Resize nếu quá lớn (giảm RAM usage và tăng tốc độ inference)
                if img_obj.size[0] > max_size[0] or img_obj.size[1] > max_size[1]:
                    img_obj.thumbnail(max_size, Image.Resampling.LANCZOS)
                
                images.append(img_obj)
                # Ước tính page number (1-indexed)
                estimated_page = max(1, (current_paragraph_count // paragraphs_per_page) + 1)
                image_page_numbers.append(estimated_page)
                
                # Clear img_data ngay sau khi tạo Image object
                del img_data
        
        # Tạo page_ranges cho text (ước tính)
        full_text = "\n".join(texts)
        page_ranges: Dict[int, Tuple[int, int]] = {}
        if texts:
            chars_per_page = len(full_text) // max(1, (len(texts) // paragraphs_per_page) + 1)
            current_pos = 0
            page_num = 1
            
            for i, text in enumerate(texts):
                if i > 0 and i % paragraphs_per_page == 0:
                    # Bắt đầu trang mới
                    page_ranges[page_num] = (current_pos, current_pos)
                    page_num += 1
                    current_pos = len("\n".join(texts[:i]))
            
            # Trang cuối
            if page_num not in page_ranges:
                page_ranges[page_num] = (current_pos, len(full_text))
        
        # Clear texts list sau khi join
        del texts
        return full_text, images, image_page_numbers, page_ranges

    # --------------------- Public API ---------------------

    async def process_pdf(self, pdf_path: str, use_parallel: bool = True) -> Tuple[str, str, List[Image.Image], List[int], Dict[int, Tuple[int, int]]]:
        """
        Xử lý PDF với tối ưu tốc độ theo từng trang:
        - Extract text trực tiếp từ PDF (nhanh, ~1-2s cho file nhiều trang)
        - Kiểm tra từng trang: nếu trang có text trực tiếp < threshold → OCR trang đó
        - Merge kết quả: dùng text trực tiếp nếu đủ tốt, nếu không thì dùng OCR (có thể lấy text từ ảnh)
        - Xử lý case text trong ảnh: nếu OCR được nhiều text hơn → dùng OCR
        
        Args:
            pdf_path: Đường dẫn file PDF
            use_parallel: Nếu True, dùng parallel OCR khi cần (nhanh hơn). Nếu False, dùng tuần tự.
        
        Returns: (full_text, text_path, images, image_page_numbers, page_ranges)
        """
        import logging
        logger = logging.getLogger(__name__)
        
        # Extract images và text trực tiếp từ PDF (song song để tăng tốc)
        (images, image_page_numbers), (direct_text, page_ranges) = await asyncio.gather(
            asyncio.to_thread(self._extract_images_from_pdf, pdf_path),
            asyncio.to_thread(self._extract_text_by_pages_from_pdf, pdf_path),
        )
        
        # Kiểm tra từng trang và quyết định trang nào cần OCR
        min_text_length = 50  # Tối thiểu 50 ký tự để coi là có text đủ tốt
        pages_to_ocr: List[int] = []
        page_direct_texts: Dict[int, str] = {}  # Lưu text trực tiếp của từng trang
        
        # Phân tích từng trang
        for page_num in sorted(page_ranges.keys()):
            start_pos, end_pos = page_ranges[page_num]
            page_direct_text = direct_text[start_pos:end_pos].strip()
            page_direct_texts[page_num] = page_direct_text
            
            # Nếu trang có ít text → cần OCR (có thể có text trong ảnh)
            if len(page_direct_text) < min_text_length:
                pages_to_ocr.append(page_num)
        
        total_pages = len(page_ranges)
        pages_needing_ocr = len(pages_to_ocr)
        
        logger.info(
            f"PDF text extraction: {total_pages - pages_needing_ocr}/{total_pages} pages have sufficient direct text. "
            f"Pages needing OCR: {pages_needing_ocr}"
        )
        
        # Nếu không có trang nào cần OCR → dùng text trực tiếp
        if not pages_to_ocr:
            logger.info(f"All pages have sufficient direct text. Using direct text extraction (fast path).")
            with tempfile.NamedTemporaryFile(mode='w', delete=False, suffix=".txt", encoding="utf-8") as tf:
                tf.write(direct_text)
                text_path = tf.name
            
            return direct_text, text_path, images, image_page_numbers, page_ranges
        
        # OCR những trang thiếu text (song song để tăng tốc)
        logger.info(f"OCR-ing {pages_needing_ocr} pages that need OCR: {pages_to_ocr}")
        
        # OCR từng trang (song song nếu use_parallel=True, tuần tự nếu False)
        max_workers = 4 if use_parallel else 1
        ocr_texts_by_page = await run_ocr_on_pages(pdf_path, pages_to_ocr, max_workers=max_workers)
        
        # Merge kết quả: ghép text từng trang theo thứ tự
        merged_page_texts: List[str] = []
        new_page_ranges: Dict[int, Tuple[int, int]] = {}
        current_pos = 0
        
        for page_num in sorted(page_ranges.keys()):
            if page_num in pages_to_ocr:
                # Trang này đã OCR → dùng OCR text
                ocr_text = ocr_texts_by_page.get(page_num, "")
                # So sánh với text trực tiếp: nếu OCR có nhiều text hơn → dùng OCR
                direct_text_page = page_direct_texts.get(page_num, "")
                if len(ocr_text.strip()) > len(direct_text_page.strip()):
                    page_text = ocr_text
                    logger.debug(f"Page {page_num}: Using OCR text ({len(ocr_text)} chars) over direct text ({len(direct_text_page)} chars)")
                else:
                    # OCR không tốt hơn → dùng text trực tiếp (nếu có)
                    page_text = direct_text_page if direct_text_page else ocr_text
                    logger.debug(f"Page {page_num}: Using direct text ({len(direct_text_page)} chars) over OCR text ({len(ocr_text)} chars)")
            else:
                # Trang này có text trực tiếp đủ tốt → dùng text trực tiếp
                page_text = page_direct_texts.get(page_num, "")
            
            # Thêm newline giữa các trang (trừ trang cuối)
            if page_num < total_pages:
                page_text += "\n"
            
            start_pos = current_pos
            end_pos = current_pos + len(page_text)
            new_page_ranges[page_num] = (start_pos, end_pos)
            merged_page_texts.append(page_text)
            current_pos = end_pos
        
        # Ghép toàn bộ text
        full_text = "".join(merged_page_texts)
        
        # Lưu vào file tạm
        with tempfile.NamedTemporaryFile(mode='w', delete=False, suffix=".txt", encoding="utf-8") as tf:
            tf.write(full_text)
            text_path = tf.name
        
        logger.info(
            f"PDF processing completed: {total_pages - pages_needing_ocr} pages used direct text, "
            f"{pages_needing_ocr} pages used OCR. Total text length: {len(full_text)} chars"
        )
        
        return full_text, text_path, images, image_page_numbers, new_page_ranges

    def process_docx(self, docx_path: str) -> Tuple[str, str, List[Image.Image], List[int], Dict[int, Tuple[int, int]]]:
        """
        Xử lý DOCX:
        - Trích text + ảnh bằng python-docx.
        - Ghi text ra file txt tạm -> text_path.
        - Trả về page numbers (ước tính) cho images và page_ranges cho text.
        
        Returns: (full_text, text_path, images, image_page_numbers, page_ranges)
        """
        full_text, images, image_page_numbers, page_ranges = self._extract_text_and_images_from_docx(docx_path)

        with tempfile.NamedTemporaryFile(delete=False, suffix=".txt") as tf:
            tf.write(full_text.encode("utf-8"))
            text_path = tf.name

        return full_text, text_path, images, image_page_numbers, page_ranges


